from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from assembler_html import AssetRef, MARKER_RE, find_transcript_files, load_assets


SECTION_HEADING_RE = re.compile(r"^###\s+(Full Width|Left Column|Right Column)\s*$")
PAGE_HEADING_RE = re.compile(r"^##\s+Page\s+(\d+)\s*$", re.IGNORECASE)
QUESTION_START_RE = re.compile(r"^\s*(\d{1,2})\s*\.")
LABEL_TEXT_RE = re.compile(r"(그림|표)\s*(\d+)")
ASSET_TEXT_REGIONS = {"table", "image", "asset"}
SUPPRESSED_LINE_REGIONS = ASSET_TEXT_REGIONS | {"header_footer"}


@dataclass
class PageContent:
    page_no: int
    full_width: list[str] = field(default_factory=list)
    left_column: list[str] = field(default_factory=list)
    right_column: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class LayoutPage:
    page_no: int
    page_width_pt: float
    page_height_pt: float
    is_two_column: bool


@dataclass(frozen=True)
class ContentBlock:
    kind: str
    text: str = ""
    asset: AssetRef | None = None


@dataclass
class QuestionContent:
    number: int
    lines: list[str] = field(default_factory=list)


@dataclass
class LineHintBundle:
    asset_text_by_label: dict[str, list[str]] = field(default_factory=dict)
    suppress_counts_by_page: dict[int, Counter[str]] = field(default_factory=dict)

    @property
    def has_hints(self) -> bool:
        return bool(self.asset_text_by_label or self.suppress_counts_by_page)


def assemble_docx(result_dir: Path, layout_mode: str = "linear", asset_text_mode: str = "below") -> Path:
    """Create an editable DOCX review file from one transcript result folder."""
    result_dir = Path(result_dir).expanduser().resolve()
    files = find_transcript_files(result_dir)
    markdown_text = files.markdown.read_text(encoding="utf-8-sig")
    assets = load_assets(result_dir)
    pages = parse_markdown_pages(markdown_text)
    line_hints = load_line_hints(result_dir, files.markdown.stem, assets)
    if line_hints.has_hints:
        pages = apply_line_hints_to_pages(pages, line_hints)
    layout_pages = load_layout_pages(result_dir, files.markdown.stem)

    output_path = _output_path(result_dir, files.markdown.stem, layout_mode)
    document = build_docx(
        markdown_stem=files.markdown.stem,
        pages=pages,
        assets=assets,
        line_hints=line_hints,
        layout_pages=layout_pages,
        layout_mode=layout_mode,
        asset_text_mode=asset_text_mode,
    )
    document.save(output_path)
    return output_path


def parse_markdown_pages(markdown_text: str) -> list[PageContent]:
    """Use markdown scaffold headings as layout hints without copying them into DOCX."""
    pages: list[PageContent] = []
    current_page: PageContent | None = None
    current_area = "full_width"

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        page_match = PAGE_HEADING_RE.match(stripped)
        if page_match:
            current_page = PageContent(page_no=int(page_match.group(1)))
            pages.append(current_page)
            current_area = "full_width"
            continue

        section_match = SECTION_HEADING_RE.match(stripped)
        if section_match:
            section_name = section_match.group(1)
            if section_name == "Left Column":
                current_area = "left_column"
            elif section_name == "Right Column":
                current_area = "right_column"
            else:
                current_area = "full_width"
            continue

        if current_page is None:
            continue
        if stripped.startswith("# ") or stripped.startswith("> "):
            continue

        getattr(current_page, current_area).append(line)

    return pages


def load_layout_pages(result_dir: Path, stem: str) -> dict[int, LayoutPage]:
    layout_path = result_dir / f"{stem}.layout.json"
    if not layout_path.exists():
        return {}

    with layout_path.open("r", encoding="utf-8-sig") as f:
        data = json.load(f)

    raw_pages = data.get("pages", []) if isinstance(data, dict) else data
    pages: dict[int, LayoutPage] = {}
    for item in raw_pages:
        if not isinstance(item, dict):
            continue
        page_no = _as_int(item.get("page_no"))
        width = _as_float(item.get("page_width"))
        height = _as_float(item.get("page_height"))
        if page_no is None or width is None or height is None:
            continue
        pages[page_no] = LayoutPage(
            page_no=page_no,
            page_width_pt=width,
            page_height_pt=height,
            is_two_column=bool(item.get("is_two_column")),
        )
    return pages


def load_line_hints(result_dir: Path, stem: str, assets: Iterable[AssetRef]) -> LineHintBundle:
    """Load optional line-level hints exported by the transcript step."""
    lines_path = result_dir / f"{stem}.lines.json"
    if not lines_path.exists():
        return LineHintBundle()

    with lines_path.open("r", encoding="utf-8-sig") as f:
        data = json.load(f)

    if isinstance(data, dict):
        raw_lines = data.get("lines") or data.get("items") or []
    elif isinstance(data, list):
        raw_lines = data
    else:
        raise ValueError(f"지원하지 않는 lines.json 형식입니다: {lines_path}")

    assets_by_id = {asset.asset_id: asset.label for asset in assets if asset.asset_id and asset.label}
    asset_text_by_label: dict[str, list[str]] = defaultdict(list)
    suppress_counts_by_page: dict[int, Counter[str]] = defaultdict(Counter)

    for item in raw_lines:
        if not isinstance(item, dict):
            continue

        text = str(item.get("text") or "").strip()
        if not text:
            continue

        region = str(item.get("region") or "").strip().lower()
        page_no = _as_int(item.get("page_no")) or 0
        asset_label = _normalize_asset_label(item.get("asset_label"))
        if not asset_label:
            asset_id = str(item.get("asset_id") or "").strip()
            asset_label = assets_by_id.get(asset_id, "")

        if region in ASSET_TEXT_REGIONS and asset_label:
            asset_text_by_label[asset_label].append(text)

        if region in SUPPRESSED_LINE_REGIONS:
            normalized = _normalize_hint_text(text)
            if normalized:
                suppress_counts_by_page[page_no][normalized] += 1

    return LineHintBundle(
        asset_text_by_label={key: _clean_asset_note_lines(value) for key, value in asset_text_by_label.items()},
        suppress_counts_by_page={key: value for key, value in suppress_counts_by_page.items()},
    )


def apply_line_hints_to_pages(pages: list[PageContent], line_hints: LineHintBundle) -> list[PageContent]:
    """Remove text that the transcript step marked as asset/header/footer content."""
    if not line_hints.has_hints:
        return pages

    filtered_pages: list[PageContent] = []
    for page in pages:
        suppress_counts = Counter(line_hints.suppress_counts_by_page.get(page.page_no, Counter()))
        suppress_counts.update(line_hints.suppress_counts_by_page.get(0, Counter()))

        filtered_pages.append(
            PageContent(
                page_no=page.page_no,
                full_width=_filter_lines_with_hints(page.full_width, suppress_counts),
                left_column=_filter_lines_with_hints(page.left_column, suppress_counts),
                right_column=_filter_lines_with_hints(page.right_column, suppress_counts),
            )
        )

    return filtered_pages


def build_docx(
    markdown_stem: str,
    pages: list[PageContent],
    assets: Iterable[AssetRef],
    line_hints: LineHintBundle | None,
    layout_pages: dict[int, LayoutPage],
    layout_mode: str = "linear",
    asset_text_mode: str = "below",
) -> Document:
    document = Document()
    default_layout = _choose_default_layout(layout_pages, layout_mode)
    _configure_document(document, default_layout)
    _configure_styles(document)

    assets_by_label = {asset.label: asset for asset in assets if asset.label}
    asset_text_by_label = _select_asset_text_by_label(line_hints, assets_by_label, asset_text_mode)
    page_width_pt = default_layout.page_width_pt
    page_height_pt = default_layout.page_height_pt
    margin_pt = _margin_pt(page_width_pt, page_height_pt)
    full_width_pt = page_width_pt - margin_pt * 2
    gutter_pt = 18.0
    column_width_pt = max((full_width_pt - gutter_pt) / 2, 180.0)

    if layout_mode == "questions":
        intro_lines, questions = split_questions(pages)
        _build_question_pages(
            document,
            markdown_stem,
            intro_lines,
            questions,
            assets_by_label,
            asset_text_by_label,
            full_width_pt,
        )
        return document

    for page_index, page in enumerate(pages):
        if page_index:
            if layout_mode == "linear":
                document.add_paragraph()
            else:
                document.add_page_break()

        _add_lines(document, page.full_width, assets_by_label, asset_text_by_label, full_width_pt)

        if layout_mode == "linear":
            _add_lines(document, page.left_column, assets_by_label, asset_text_by_label, full_width_pt)
            _add_lines(document, page.right_column, assets_by_label, asset_text_by_label, full_width_pt)
            continue

        if page.left_column or page.right_column:
            left_blocks = _lines_to_blocks(page.left_column, assets_by_label)
            right_blocks = _lines_to_blocks(page.right_column, assets_by_label)
            row_count = max(len(left_blocks), len(right_blocks), 1)
            table = document.add_table(rows=row_count, cols=3)
            table.autofit = False
            _set_table_fixed_layout(table)
            _set_table_grid(table, [column_width_pt, gutter_pt, column_width_pt])
            _set_table_borders(table, "single", "FFFFFF")
            _set_table_cell_margins(table, top=0, start=0, bottom=0, end=0)

            for row_index, row in enumerate(table.rows):
                left_cell, gutter_cell, right_cell = row.cells
                for cell, width in ((left_cell, column_width_pt), (gutter_cell, gutter_pt), (right_cell, column_width_pt)):
                    _set_cell_width(cell, width)
                    _set_cell_borders(cell, "single", "FFFFFF")
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    _clear_cell(cell)
                if row_index < len(left_blocks):
                    _add_block(left_cell, left_blocks[row_index], column_width_pt, asset_text_by_label)
                if row_index < len(right_blocks):
                    _add_block(right_cell, right_blocks[row_index], column_width_pt, asset_text_by_label)

    if not pages:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{markdown_stem}: 조립할 페이지 본문이 없습니다.")

    return document


def split_questions(pages: list[PageContent]) -> tuple[list[str], list[QuestionContent]]:
    intro_lines: list[str] = []
    questions: list[QuestionContent] = []
    current_question: QuestionContent | None = None

    if pages:
        intro_lines.extend(_filter_question_lines(pages[0].full_width))

    for page in pages:
        for line in _filter_question_lines([*page.left_column, *page.right_column]):
            match = QUESTION_START_RE.match(line.strip())
            if match:
                current_question = QuestionContent(number=int(match.group(1)))
                questions.append(current_question)
            if current_question is None:
                intro_lines.append(line)
            else:
                current_question.lines.append(line)

    return intro_lines, questions


def _filter_question_lines(lines: list[str]) -> list[str]:
    filtered: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            filtered.append(line)
            continue
        if stripped == "---":
            continue
        if re.fullmatch(r"\d{1,2}", stripped):
            continue
        if stripped in {"생활과 윤리", "사회탐구", "영역"}:
            continue
        filtered.append(line)
    return filtered


def _build_question_pages(
    document: Document,
    markdown_stem: str,
    intro_lines: list[str],
    questions: list[QuestionContent],
    assets_by_label: dict[str, AssetRef],
    asset_text_by_label: dict[str, list[str]],
    full_width_pt: float,
) -> None:
    _add_review_title(document, markdown_stem, len(questions))
    if intro_lines:
        _add_section_label(document, "시험지 정보")
        _add_lines(document, intro_lines, assets_by_label, asset_text_by_label, full_width_pt, font_size=10.0)

    for index, question in enumerate(questions):
        if index or intro_lines:
            document.add_page_break()
        _add_question_heading(document, question.number)
        _add_lines(document, question.lines, assets_by_label, asset_text_by_label, full_width_pt, font_size=10.5)

    if not questions:
        _add_text_paragraph(document, "분리할 문항을 찾지 못했습니다.", font_size=10.5)


def _add_review_title(document: Document, markdown_stem: str, question_count: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(markdown_stem)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run(f"문항별 검수본 · {question_count}문항")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(80, 96, 112)
    meta_run.font.name = "Malgun Gothic"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _add_section_label(document: Document, label: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 122, 104)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _add_question_heading(document: Document, number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(f"문항 {number}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 122, 104)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _configure_document(document: Document, layout: LayoutPage) -> None:
    section = document.sections[0]
    _configure_section(section, layout)
    _set_section_columns(section, 1)


def _configure_section(section, layout: LayoutPage) -> None:
    section.page_width = Pt(layout.page_width_pt)
    section.page_height = Pt(layout.page_height_pt)
    margin = _margin_pt(layout.page_width_pt, layout.page_height_pt)
    section.top_margin = Pt(margin)
    section.bottom_margin = Pt(margin)
    section.left_margin = Pt(margin)
    section.right_margin = Pt(margin)
    section.header_distance = Pt(18)
    section.footer_distance = Pt(18)


def _set_section_columns(section, count: int, space_pt: float = 18.0) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(int(space_pt * 20)))


def _add_column_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_break(WD_BREAK.COLUMN)


def _output_path(result_dir: Path, stem: str, layout_mode: str) -> Path:
    if layout_mode == "questions":
        return result_dir / f"{stem}.assembled.questions.docx"
    if layout_mode == "table":
        return result_dir / f"{stem}.assembled.table.docx"
    return result_dir / f"{stem}.assembled.docx"


def _lines_to_blocks(lines: list[str], assets_by_label: dict[str, AssetRef]) -> list[ContentBlock]:
    blocks: list[ContentBlock] = []
    previous_blank = True
    for line in lines:
        if not line.strip():
            if not previous_blank:
                blocks.append(ContentBlock(kind="blank"))
            previous_blank = True
            continue
        blocks.extend(_line_to_blocks(line, assets_by_label))
        previous_blank = False
    return blocks


def _line_to_blocks(line: str, assets_by_label: dict[str, AssetRef]) -> list[ContentBlock]:
    blocks: list[ContentBlock] = []
    position = 0

    for match in MARKER_RE.finditer(line):
        if match.start() > position:
            text = line[position : match.start()].strip()
            if text:
                blocks.append(ContentBlock(kind="text", text=text))

        label = f"{match.group(1)} {int(match.group(2))}"
        asset = assets_by_label.get(label)
        if asset is None or not asset.exists:
            detail = asset.raw_file if asset is not None else "assets.json에 해당 항목이 없습니다."
            blocks.append(ContentBlock(kind="missing", text=f"{label}|{detail or ''}"))
        else:
            blocks.append(ContentBlock(kind="asset", asset=asset))
        position = match.end()

    if position < len(line):
        tail = line[position:].strip()
        if tail:
            blocks.append(ContentBlock(kind="text", text=tail))

    return blocks or [ContentBlock(kind="text", text=line.strip())]


def _add_block(
    container,
    block: ContentBlock,
    max_width_pt: float,
    asset_text_by_label: dict[str, list[str]] | None = None,
) -> None:
    if block.kind == "blank":
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
    elif block.kind == "asset" and block.asset is not None:
        _add_asset_picture(container, block.asset, max_width_pt, _asset_notes(block.asset, asset_text_by_label))
    elif block.kind == "missing":
        label, _, detail = block.text.partition("|")
        _add_missing_paragraph(container, label, detail)
    else:
        _add_text_paragraph(container, block.text)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(8.7)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    for style_name in ("Body Text", "Caption"):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Malgun Gothic"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _add_lines(
    container,
    lines: list[str],
    assets_by_label: dict[str, AssetRef],
    asset_text_by_label: dict[str, list[str]] | None,
    max_width_pt: float,
    font_size: float = 8.7,
) -> None:
    previous_blank = True
    for line in lines:
        if not line.strip():
            if not previous_blank:
                paragraph = container.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
            previous_blank = True
            continue

        _add_line(container, line, assets_by_label, asset_text_by_label, max_width_pt, font_size=font_size)
        previous_blank = False


def _add_line(
    container,
    line: str,
    assets_by_label: dict[str, AssetRef],
    asset_text_by_label: dict[str, list[str]] | None,
    max_width_pt: float,
    font_size: float = 8.7,
) -> None:
    position = 0
    text_buffer: list[str] = []

    for match in MARKER_RE.finditer(line):
        if match.start() > position:
            text_buffer.append(line[position : match.start()])

        buffered_text = "".join(text_buffer).strip()
        if buffered_text:
            _add_text_paragraph(container, buffered_text, font_size=font_size)
            text_buffer.clear()

        label = f"{match.group(1)} {int(match.group(2))}"
        asset = assets_by_label.get(label)
        if asset is None or not asset.exists:
            detail = asset.raw_file if asset is not None else "assets.json에 해당 항목이 없습니다."
            _add_missing_paragraph(container, label, detail)
        else:
            _add_asset_picture(container, asset, max_width_pt, _asset_notes(asset, asset_text_by_label))
        position = match.end()

    if position < len(line):
        text_buffer.append(line[position:])

    tail_text = "".join(text_buffer).strip()
    if tail_text:
        _add_text_paragraph(container, tail_text, font_size=font_size)


def _add_text_paragraph(container, text: str, font_size: float = 8.7):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(font_size)
    return paragraph


def _add_missing_paragraph(container, label: str, detail: str | None) -> None:
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"파일 없음: {label}")
    run.bold = True
    run.font.color.rgb = RGBColor(155, 93, 0)
    if detail:
        paragraph.add_run(f" ({detail})")


def _add_asset_picture(
    container,
    asset: AssetRef,
    max_width_pt: float,
    extracted_lines: list[str] | None = None,
) -> None:
    caption = container.add_paragraph()
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(1)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(f"[{asset.label}]")
    caption_run.bold = True
    caption_run.font.size = Pt(7.5)
    caption_run.font.color.rgb = RGBColor(31, 122, 104)
    caption_run.font.name = "Malgun Gothic"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    image_width_pt = _fit_asset_width(asset, max_width_pt)
    run.add_picture(str(asset.resolved_file), width=Pt(image_width_pt))
    if extracted_lines:
        _add_asset_text_note(container, extracted_lines)


def _add_asset_text_note(container, lines: list[str]) -> None:
    note_lines = _clean_asset_note_lines(lines)
    if not note_lines:
        return

    label = container.add_paragraph()
    label.paragraph_format.space_before = Pt(0)
    label.paragraph_format.space_after = Pt(1)
    label.paragraph_format.left_indent = Pt(10)
    run = label.add_run("추출 텍스트")
    run.bold = True
    run.font.size = Pt(8.0)
    run.font.color.rgb = RGBColor(80, 96, 112)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    for line in note_lines:
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0.6)
        paragraph.paragraph_format.left_indent = Pt(14)
        paragraph.paragraph_format.line_spacing = 1.0
        note_run = paragraph.add_run(line)
        note_run.font.size = Pt(8.0)
        note_run.font.color.rgb = RGBColor(80, 96, 112)
        note_run.font.name = "Malgun Gothic"
        note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _fit_asset_width(asset: AssetRef, max_width_pt: float) -> float:
    if asset.width and asset.width > 0:
        return min(asset.width, max_width_pt)
    return max_width_pt


def _asset_notes(asset: AssetRef, asset_text_by_label: dict[str, list[str]] | None) -> list[str]:
    if not asset_text_by_label:
        return []
    return asset_text_by_label.get(asset.label, [])


def _select_asset_text_by_label(
    line_hints: LineHintBundle | None,
    assets_by_label: dict[str, AssetRef],
    asset_text_mode: str,
) -> dict[str, list[str]]:
    if not line_hints or asset_text_mode == "none":
        return {}
    if asset_text_mode == "tables_only":
        return {
            label: lines
            for label, lines in line_hints.asset_text_by_label.items()
            if (assets_by_label.get(label) is not None and assets_by_label[label].kind == "table")
        }
    return line_hints.asset_text_by_label


def _filter_lines_with_hints(lines: list[str], suppress_counts: Counter[str]) -> list[str]:
    filtered: list[str] = []
    for line in lines:
        text_without_markers = MARKER_RE.sub("", line).strip()
        if not text_without_markers:
            filtered.append(line)
            continue

        normalized = _normalize_hint_text(text_without_markers)
        if normalized and suppress_counts[normalized] > 0:
            suppress_counts[normalized] -= 1
            marker_text = " ".join(match.group(0) for match in MARKER_RE.finditer(line))
            if marker_text:
                filtered.append(marker_text)
            continue

        filtered.append(line)
    return filtered


def _clean_asset_note_lines(lines: Iterable[str]) -> list[str]:
    cleaned: list[str] = []
    seen_consecutive = ""
    for line in lines:
        stripped = str(line).strip()
        if not stripped or stripped == "---":
            continue
        normalized = _normalize_hint_text(stripped)
        if not normalized or normalized == seen_consecutive:
            continue
        cleaned.append(stripped)
        seen_consecutive = normalized
    return cleaned


def _normalize_asset_label(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip().strip("[]")
    match = LABEL_TEXT_RE.search(text)
    if not match:
        return text
    return f"{match.group(1)} {int(match.group(2))}"


def _normalize_hint_text(text: str) -> str:
    text = str(text).strip()
    if not text:
        return ""
    text = MARKER_RE.sub("", text)
    return re.sub(r"\s+", "", text)


def _choose_default_layout(layout_pages: dict[int, LayoutPage], layout_mode: str = "linear") -> LayoutPage:
    if layout_mode == "questions":
        return LayoutPage(page_no=1, page_width_pt=595.276, page_height_pt=841.89, is_two_column=False)
    if layout_pages:
        return layout_pages[sorted(layout_pages)[0]]
    return LayoutPage(page_no=1, page_width_pt=841.0, page_height_pt=1190.0, is_two_column=True)


def _margin_pt(page_width_pt: float, page_height_pt: float) -> float:
    shortest = min(page_width_pt, page_height_pt)
    if shortest >= 800:
        return 36.0
    return 28.0


def _set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_borders(table, value: str, color: str = "auto") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), value)
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _set_table_grid(table, widths_pt: list[float]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width_pt in widths_pt:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_pt * 20)))
        grid.append(col)


def _set_table_cell_margins(table, top: int, start: int, bottom: int, end: int) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)

    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_pt: float) -> None:
    cell.width = Pt(width_pt)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_pt * 20)))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_borders(cell, value: str, color: str = "auto") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), value)
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _clear_cell(cell) -> None:
    cell._tc.clear_content()


def _as_int(value: object) -> int | None:
    try:
        return int(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _as_float(value: object) -> float | None:
    try:
        return float(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _configure_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if reconfigure is not None:
            try:
                reconfigure(encoding="utf-8")
            except Exception:
                pass


def main(argv: list[str] | None = None) -> int:
    _configure_stdio()
    parser = argparse.ArgumentParser(description="전사 결과 폴더의 텍스트와 에셋을 편집 가능한 DOCX로 조립합니다.")
    parser.add_argument("result_dir", type=Path, help="전사 결과 폴더 경로")
    parser.add_argument(
        "--layout",
        choices=("linear", "table", "questions"),
        default="linear",
        help="DOCX 배치 방식. questions는 문항별 검수본을 만듭니다.",
    )
    parser.add_argument(
        "--asset-text",
        choices=("tables_only", "below", "none"),
        default="tables_only",
        help="lines.json의 표/그림 내부 텍스트 표시 방식입니다. 기본값은 표만 보조 표시합니다.",
    )
    args = parser.parse_args(argv)

    try:
        output_path = assemble_docx(args.result_dir, layout_mode=args.layout, asset_text_mode=args.asset_text)
    except Exception as exc:
        print(f"오류: {exc}", file=sys.stderr)
        return 1

    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
